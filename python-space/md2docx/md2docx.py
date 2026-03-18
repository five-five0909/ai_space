"""
Markdown to Word Document Converter (增强版)
支持图片、表格、公式等元素,符合指定的文档格式规范
增强功能:
- 表格跨页自动添加"续表"标题
- 图片自动调整大小避免超出页面
- 支持相对路径的图片/视频
- 大标题后自动分页
"""

import re
import os
from urllib.parse import unquote
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO
import base64
from PIL import Image

# 导入公式渲染库
try:
    from math2docx import add_math
    HAS_MATH2DOCX = True
except ImportError:
    HAS_MATH2DOCX = False
    print("警告: math2docx 未安装,公式将以文本形式显示")


class MarkdownToDocxConverter:
    """Markdown转Word文档转换器"""
    
    def __init__(self, input_file_path=None):
        self.doc = Document()
        self.setup_document_style()
        self.image_counter = 0
        self.table_counter = 0
        self.current_list_level = 0
        self.input_file_dir = os.path.dirname(os.path.abspath(input_file_path)) if input_file_path else os.getcwd()
        
        # 页面尺寸常量 (用于计算是否需要分页)
        self.page_width = 21  # cm
        self.page_height = 29.7  # cm
        self.usable_height = 24  # cm (去掉页边距后的可用高度)
        self.usable_width = 15  # cm (去掉页边距后的可用宽度)
        
    def setup_document_style(self):
        """设置文档基本样式"""
        # 设置页面为A4
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        
        # 设置默认字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal'].font.size = Pt(12)  # 小四
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置段落格式
        paragraph_format = self.doc.styles['Normal'].paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(20)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Cm(0.74)  # 2字符
        
    def set_cell_border(self, cell):
        """设置表格单元格边框"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # 创建边框元素
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    def set_picture_wrap_top_bottom(self, inline_shape):
        """设置图片为上下型环绕（Top and Bottom wrap）
        
        通过将 inline 图片转换为 anchor 图片并设置 wrapTopAndBottom 来实现
        """
        from lxml import etree
        
        # 获取 inline 元素
        inline = inline_shape._inline
        
        # 获取尺寸
        cx = inline.extent.cx
        cy = inline.extent.cy
        
        # 定义命名空间
        nsmap = {
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
        }
        
        # 创建 anchor 元素
        anchor = etree.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor', nsmap=nsmap)
        
        # 设置 anchor 属性
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251658240')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        
        # simplePos 元素
        simplePos = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        
        # positionH - 水平位置相对于列居中
        positionH = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH')
        positionH.set('relativeFrom', 'column')
        align_h = etree.SubElement(positionH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
        align_h.text = 'center'
        
        # positionV - 垂直位置相对于段落
        positionV = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV')
        positionV.set('relativeFrom', 'paragraph')
        posOffset = etree.SubElement(positionV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset')
        posOffset.text = '0'
        
        # extent - 尺寸
        extent = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
        extent.set('cx', str(cx))
        extent.set('cy', str(cy))
        
        # effectExtent
        effectExtent = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}effectExtent')
        effectExtent.set('l', '0')
        effectExtent.set('t', '0')
        effectExtent.set('r', '0')
        effectExtent.set('b', '0')
        
        # wrapTopAndBottom - 上下型环绕
        wrapTopAndBottom = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}wrapTopAndBottom')
        
        # docPr - 文档属性
        docPr = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')
        docPr.set('id', str(self.image_counter))
        docPr.set('name', f'Picture {self.image_counter}')
        
        # cNvGraphicFramePr
        cNvGraphicFramePr = etree.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr')
        graphicFrameLocks = etree.SubElement(cNvGraphicFramePr, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicFrameLocks')
        graphicFrameLocks.set('noChangeAspect', '1')
        
        # 复制 graphic 元素
        graphic = inline.graphic
        anchor.append(graphic)
        
        # 替换 inline 为 anchor
        drawing = inline.getparent()
        drawing.remove(inline)
        drawing.append(anchor)
        
    def is_chinese_char(self, char):
        """判断字符是否为中文"""
        return '\u4e00' <= char <= '\u9fff'
        
    def set_run_font(self, run, text):
        """根据文本内容设置字体"""
        # 检查是否包含中文
        has_chinese = any(self.is_chinese_char(c) for c in text)
        
        if has_chinese:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            
    def add_heading(self, text, level):
        """添加标题 - 使用 Word 内置标题样式以支持导航窗格"""
        # 清除 markdown 格式符号
        clean_text = self._strip_markdown_formatting(text)
        
        # 检查标题中是否包含公式
        has_formula = '$' in text
        
        # 使用 Word 内置的 Heading 样式（level 1-9）
        heading_level = min(level, 9)
        
        if has_formula:
            # 如果包含公式，需要手动创建标题并解析公式
            paragraph = self.doc.add_paragraph()
            # 设置标题样式
            paragraph.style = f'Heading {heading_level}'
            # 解析标题内容（包含公式）
            self._parse_heading_content(paragraph, text, level)
        else:
            # 普通标题直接创建
            paragraph = self.doc.add_heading(clean_text, level=heading_level)
            # 设置字体格式
            self._set_heading_font(paragraph, level)
        
        # 设置段落格式
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _strip_markdown_formatting(self, text):
        """清除 markdown 格式符号"""
        # 清除粗斜体 ***text*** 或 ___text___
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'___(.+?)___', r'\1', text)
        # 清除粗体 **text** 或 __text__
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        # 清除斜体 *text* 或 _text_
        text = re.sub(r'\*([^*]+?)\*', r'\1', text)
        text = re.sub(r'_([^_]+?)_', r'\1', text)
        # 清除行内代码 `text`
        text = re.sub(r'`([^`]+?)`', r'\1', text)
        # 清除删除线 ~~text~~
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        return text
    
    def _parse_heading_content(self, paragraph, text, level):
        """解析标题内容，处理公式和格式"""
        # 先清除格式符号
        text = self._strip_markdown_formatting(text)
        
        # 分割公式和文本
        parts = re.split(r'(\$[^$]+\$)', text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('$') and part.endswith('$'):
                # 这是公式
                formula = part[1:-1]  # 去掉 $ 符号
                self._add_latex_formula(paragraph, formula)
            else:
                # 普通文本
                run = paragraph.add_run(part)
                self._set_heading_run_font(run, level)
    
    def _set_heading_font(self, paragraph, level):
        """设置标题段落的字体格式"""
        for run in paragraph.runs:
            self._set_heading_run_font(run, level)
    
    def _set_heading_run_font(self, run, level):
        """设置标题 run 的字体格式"""
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        if level == 1:  # 章标题 - 三号黑体
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 2:  # 节标题 - 小三黑体
            run.font.size = Pt(15)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 3:  # 条标题 - 四号黑体
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        else:  # 更低级别标题 - 小四黑体
            run.font.size = Pt(12)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
    def add_paragraph_with_format(self, text, style='Normal', **kwargs):
        """添加格式化段落"""
        paragraph = self.doc.add_paragraph(style=style)
        
        # 设置段落格式
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        if kwargs.get('no_indent'):
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            
        if kwargs.get('alignment'):
            paragraph.alignment = kwargs['alignment']
            
        # 处理富文本(粗体、斜体、代码等)
        self.parse_inline_formatting(paragraph, text)
        
        return paragraph
        
    def parse_inline_formatting(self, paragraph, text):
        """解析行内格式(粗体、斜体、代码、公式、图片等)"""
        # 匹配各种格式的正则表达式 - 按优先级排序
        patterns = [
            (r'!\[([^\]]*)\]\(([^)]+)\)', 'inline_image'),  # 行内图片
            (r'\$\$(.+?)\$\$', 'block_formula'),  # 块级公式
            (r'\$(.+?)\$', 'inline_formula'),  # 行内公式
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # 粗斜体
            (r'___(.+?)___', 'bold_italic'),
            (r'\*\*(.+?)\*\*', 'bold'),  # 粗体
            (r'__(.+?)__', 'bold'),
            (r'\*([^*]+?)\*', 'italic'),  # 斜体 - 修改正则避免贪婪匹配
            (r'_([^_]+?)_', 'italic'),
            (r'`([^`]+?)`', 'code'),  # 行内代码
            (r'~~(.+?)~~', 'strike'),  # 删除线
        ]
        
        # 解析文本为段落
        remaining_text = text
        segments = []
        
        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            earliest_pattern = None
            earliest_full_match = None
            
            # 找到最早出现的格式
            for pattern, format_type in patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    earliest_pattern = format_type
                    earliest_full_match = match.group(0)
                    
            if earliest_match:
                # 添加格式前的普通文本
                if earliest_pos > 0:
                    segments.append(('normal', remaining_text[:earliest_pos], None))
                
                # 处理不同类型的格式
                if earliest_pattern == 'inline_image':
                    # 行内图片 - 存储 alt_text 和 url
                    alt_text = earliest_match.group(1).strip()
                    url = earliest_match.group(2).strip()
                    segments.append(('inline_image', alt_text, url))
                elif earliest_pattern in ('inline_formula', 'block_formula'):
                    # 公式
                    formula = earliest_match.group(1)
                    segments.append((earliest_pattern, formula, None))
                else:
                    # 其他格式化文本
                    segments.append((earliest_pattern, earliest_match.group(1), None))
                    
                remaining_text = remaining_text[earliest_match.end():]
            else:
                # 没有更多格式,添加剩余文本
                if remaining_text:
                    segments.append(('normal', remaining_text, None))
                break
                
        # 创建runs
        for format_type, content, extra in segments:
            if format_type == 'inline_image':
                # 行内图片需要特殊处理 - 添加到当前段落
                self._add_inline_image(paragraph, extra, content)
            elif format_type in ('inline_formula', 'block_formula'):
                # 使用 math2docx 渲染 LaTeX 公式
                self._add_latex_formula(paragraph, content)
            else:
                run = paragraph.add_run(content)
                self.set_run_font(run, content)
                run.font.size = Pt(12)
                
                if 'bold' in format_type:
                    run.font.bold = True
                if 'italic' in format_type:
                    run.font.italic = True
                if format_type == 'code':
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(199, 37, 78)
                if format_type == 'strike':
                    run.font.strike = True
    
    def _add_latex_formula(self, paragraph, latex_code):
        """添加 LaTeX 公式到段落
        
        使用 math2docx 将 LaTeX 转换为 Office Math (OMML) 格式
        """
        if HAS_MATH2DOCX:
            try:
                # 使用 math2docx 添加公式
                add_math(paragraph, latex_code)
            except Exception as e:
                # 如果渲染失败，回退到文本显示
                run = paragraph.add_run(latex_code)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(12)
                run.font.italic = True
                print(f"警告: 公式渲染失败 '{latex_code}': {e}")
        else:
            # 没有 math2docx，使用文本显示
            run = paragraph.add_run(latex_code)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            run.font.italic = True
    
    def _add_inline_image(self, paragraph, url, alt_text=""):
        """添加行内图片到指定段落"""
        resolved_path = self.resolve_image_path(url)
        if not resolved_path:
            run = paragraph.add_run(f'[图片: {alt_text}]')
            run.font.color.rgb = RGBColor(255, 0, 0)
            return
            
        try:
            # 处理 webp 格式
            if resolved_path.lower().endswith('.webp'):
                img = Image.open(resolved_path)
                image_stream = BytesIO()
                img.save(image_stream, format='PNG')
                image_stream.seek(0)
            else:
                image_stream = resolved_path
            
            # 获取图片尺寸并自适应调整
            dimensions = self.get_image_dimensions(resolved_path)
            display_width = Inches(4)  # 默认宽度
            
            if dimensions:
                width_px, height_px = dimensions
                # 转换为英寸 (假设96 DPI)
                width_inches = width_px / 96
                # 行内图片最大宽度限制在页面可用宽度的80%
                max_width = self.usable_width * 0.8 / 2.54  # 转换为英寸
                if width_inches > max_width:
                    display_width = Inches(max_width)
                else:
                    display_width = Inches(min(width_inches, 5))
            
            # 行内图片
            self.image_counter += 1
            run = paragraph.add_run()
            picture = run.add_picture(image_stream, width=display_width)
            
            # 设置图片为上下型环绕
            self.set_picture_wrap_top_bottom(picture)
        except Exception as e:
            run = paragraph.add_run(f'[图片加载失败]')
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    def resolve_image_path(self, path):
        """解析图片路径(支持相对路径和绝对路径)"""
        # URL直接返回
        if path.startswith(('http://', 'https://', 'data:image')):
            return path
        
        # 解码 URL 编码的路径（如 %20 -> 空格）
        path = unquote(path)
        
        # 绝对路径
        if os.path.isabs(path):
            if os.path.exists(path):
                return path
            else:
                print(f"警告: 图片文件不存在: {path}")
                return None
        
        # 相对路径 - 相对于MD文件所在目录
        full_path = os.path.join(self.input_file_dir, path)
        if os.path.exists(full_path):
            return full_path
        
        # 相对于当前工作目录
        if os.path.exists(path):
            return path
            
        print(f"警告: 图片文件不存在: {path}")
        return None
    
    def get_image_dimensions(self, image_source):
        """获取图片尺寸"""
        try:
            if isinstance(image_source, str):
                if image_source.startswith(('http://', 'https://')):
                    response = requests.get(image_source, timeout=10)
                    img = Image.open(BytesIO(response.content))
                elif image_source.startswith('data:image'):
                    header, encoded = image_source.split(',', 1)
                    image_data = base64.b64decode(encoded)
                    img = Image.open(BytesIO(image_data))
                else:
                    img = Image.open(image_source)
            else:
                img = Image.open(image_source)
            
            # 返回尺寸(像素)
            return img.size
        except Exception as e:
            print(f"警告: 无法获取图片尺寸: {e}")
            return None
                
    def add_image(self, url_or_path, alt_text=""):
        """添加图片(自动调整大小)"""
        self.image_counter += 1
        
        # 解析图片路径
        resolved_path = self.resolve_image_path(url_or_path)
        if not resolved_path:
            # 图片路径无效,添加占位文本
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(f'[图片未找到: {alt_text} - 路径: {url_or_path}]')
            run.font.color.rgb = RGBColor(255, 0, 0)
            return
        
        try:
            # 获取图片
            if resolved_path.startswith(('http://', 'https://')):
                response = requests.get(resolved_path, timeout=10)
                image_stream = BytesIO(response.content)
            elif resolved_path.startswith('data:image'):
                header, encoded = resolved_path.split(',', 1)
                image_data = base64.b64decode(encoded)
                image_stream = BytesIO(image_data)
            else:
                # 本地文件 - 检查是否为 webp 格式
                if resolved_path.lower().endswith('.webp'):
                    # webp 格式需要转换为 PNG
                    img = Image.open(resolved_path)
                    image_stream = BytesIO()
                    img.save(image_stream, format='PNG')
                    image_stream.seek(0)
                else:
                    image_stream = resolved_path
            
            # 获取图片原始尺寸
            dimensions = self.get_image_dimensions(resolved_path)
            
            # 计算合适的显示宽度
            max_width_cm = self.usable_width - 2  # 留2cm边距
            max_width_inches = max_width_cm / 2.54
            
            # 默认宽度
            display_width = Inches(5)
            
            if dimensions:
                width_px, height_px = dimensions
                # 转换为英寸 (假设96 DPI)
                width_inches = width_px / 96
                height_inches = height_px / 96
                
                # 如果图片太宽,按比例缩小
                if width_inches > max_width_inches:
                    display_width = Inches(max_width_inches)
                else:
                    display_width = Inches(min(width_inches, 5))  # 最大5英寸
            
            # 添加图片
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(20)
            
            run = paragraph.add_run()
            picture = run.add_picture(image_stream, width=display_width)
            
            # 设置图片为上下型环绕
            self.set_picture_wrap_top_bottom(picture)
            
            # 添加图题
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.paragraph_format.first_line_indent = Pt(0)
            caption_run = caption_para.add_run(f'图 {self.image_counter} {alt_text}')
            caption_run.font.size = Pt(10.5)  # 五号
            self.set_run_font(caption_run, alt_text)
            
        except Exception as e:
            # 图片加载失败,添加占位文本
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(f'[图片加载失败: {alt_text} - 错误: {str(e)}]')
            run.font.color.rgb = RGBColor(255, 0, 0)
            print(f"错误: 图片加载失败 - {url_or_path}: {str(e)}")
    
    def estimate_table_rows_per_page(self):
        """估算每页能容纳的表格行数"""
        # 假设每行高度约1cm (包括内容和行距)
        return int(self.usable_height)
            
    def add_table(self, rows, is_continuation=False, original_table_num=None):
        """添加表格(支持跨页续表)"""
        if not rows:
            return
        
        # 如果不是续表,增加表格计数
        if not is_continuation:
            self.table_counter += 1
            current_table_num = self.table_counter
        else:
            current_table_num = original_table_num
        
        # 估算每页能容纳的行数
        max_rows_per_page = self.estimate_table_rows_per_page()
        
        # 如果表格行数超过一页能容纳的量,需要分页
        if len(rows) > max_rows_per_page and not is_continuation:
            # 分割表格
            header = rows[0]
            data_rows = rows[1:]
            
            # 第一部分
            first_part = [header] + data_rows[:max_rows_per_page-1]
            self.add_table(first_part, is_continuation=False)
            
            # 后续部分
            remaining = data_rows[max_rows_per_page-1:]
            while remaining:
                self.doc.add_page_break()
                chunk = remaining[:max_rows_per_page-1]
                remaining = remaining[max_rows_per_page-1:]
                continuation_rows = [header] + chunk
                self.add_table(continuation_rows, is_continuation=True, 
                             original_table_num=current_table_num)
            return
        
        # 添加表题
        caption_para = self.doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.paragraph_format.first_line_indent = Pt(0)
        
        if is_continuation:
            caption_text = f'续表 {current_table_num}'
        else:
            caption_text = f'表 {current_table_num}'
            
        caption_run = caption_para.add_run(caption_text)
        caption_run.font.size = Pt(10.5)  # 五号
        caption_run.font.name = '宋体'
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        caption_run.font.bold = True
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # 填充表格内容
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                
                # 设置单元格格式
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing = Pt(20)
                    
                    for run in paragraph.runs:
                        run.font.size = Pt(12)  # 小四
                        self.set_run_font(run, cell_text)
                        
                        # 表头加粗
                        if i == 0:
                            run.font.bold = True
                            
                # 设置边框
                self.set_cell_border(cell)
        
        # 设置表格允许跨页
        table._element.xpath('./w:tblPr')[0].append(
            OxmlElement('w:tblPrEx')
        )
                
        # 添加空行
        self.doc.add_paragraph()
        
    def add_code_block(self, code, language=''):
        """添加代码块"""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Cm(1)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(18)
        
        run = paragraph.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
        
        # 添加浅灰色背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        run._element.rPr.append(shading)
        
    def add_formula(self, latex_code):
        """添加数学公式（块级，居中显示）"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        
        # 使用 math2docx 渲染公式
        self._add_latex_formula(paragraph, latex_code)
        
    def add_list_item(self, text, ordered=False, level=0):
        """添加列表项"""
        paragraph = self.doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
        paragraph.paragraph_format.left_indent = Cm(0.74 * (level + 1))
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(20)
        
        self.parse_inline_formatting(paragraph, text)
        
    def parse_markdown_table(self, table_text):
        """解析Markdown表格"""
        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
        
        if len(lines) < 2:
            return None
            
        # 解析表头
        header = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
        
        # 跳过分隔行
        rows = [header]
        
        # 解析数据行
        for line in lines[2:]:
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                rows.append(row)
                
        return rows
        
    def convert(self, markdown_text, output_path='output.docx'):
        """转换Markdown到Word文档"""
        lines = markdown_text.split('\n')
        i = 0
        
        in_code_block = False
        code_block_content = []
        code_language = ''
        
        in_table = False
        table_content = []
        
        in_formula_block = False
        formula_content = []
        
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_language = line.strip()[3:].strip()
                    code_block_content = []
                else:
                    in_code_block = False
                    self.add_code_block('\n'.join(code_block_content), code_language)
                    code_block_content = []
                i += 1
                continue
                
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
                
            # 处理数学公式块
            if line.strip().startswith('$$'):
                if not in_formula_block:
                    in_formula_block = True
                    formula_content = []
                else:
                    in_formula_block = False
                    self.add_formula('\n'.join(formula_content))
                    formula_content = []
                i += 1
                continue
                
            if in_formula_block:
                formula_content.append(line)
                i += 1
                continue
                
            # 处理标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.add_heading(text, level)
                i += 1
                continue
                
            # 处理表格
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_content = [line]
                else:
                    table_content.append(line)
                i += 1
                
                # 检查是否表格结束
                if i >= len(lines) or '|' not in lines[i]:
                    rows = self.parse_markdown_table('\n'.join(table_content))
                    if rows:
                        self.add_table(rows)
                    in_table = False
                    table_content = []
                continue
                
            # 处理图片 - 检查是否整行只有图片
            # 匹配所有图片
            image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
            images_in_line = re.findall(image_pattern, line.strip())
            line_without_images = re.sub(image_pattern, '', line.strip())
            
            # 如果整行只有图片（去掉图片后只剩空白或零宽字符）
            if images_in_line and not line_without_images.replace('\u200b', '').replace('\u200c', '').strip():
                # 独立图片行 - 每个图片单独处理
                for alt_text, url in images_in_line:
                    alt_text = alt_text.strip()
                    url = url.strip()
                    
                    # 如果alt_text为空或者是无意义的文件名,尝试从路径中提取更好的描述
                    if not alt_text or re.match(r'^image\d+', alt_text):
                        filename = os.path.splitext(os.path.basename(url))[0]
                        cleaned_name = re.sub(r'-\d{14}-[a-z0-9]+$', '', filename)
                        alt_text = cleaned_name if cleaned_name != filename else alt_text
                    
                    self.add_image(url, alt_text)
                i += 1
                continue
                
            # 处理无序列表
            list_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                level = indent // 2
                self.add_list_item(text, ordered=False, level=level)
                i += 1
                continue
                
            # 处理有序列表
            ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if ordered_match:
                indent = len(ordered_match.group(1))
                text = ordered_match.group(3)
                level = indent // 2
                self.add_list_item(text, ordered=True, level=level)
                i += 1
                continue
                
            # 处理分隔线
            if re.match(r'^(-{3,}|\*{3,}|_{3,})$', line.strip()):
                self.doc.add_page_break()
                i += 1
                continue
                
            # 处理引用
            if line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                paragraph = self.add_paragraph_with_format(quote_text, no_indent=True)
                paragraph.paragraph_format.left_indent = Cm(1)
                paragraph.paragraph_format.right_indent = Cm(1)
                for run in paragraph.runs:
                    run.font.italic = True
                i += 1
                continue
                
            # 处理普通段落 (包含行内图片、公式等的处理在 parse_inline_formatting 中)
            if line.strip():
                self.add_paragraph_with_format(line.strip())
            else:
                # 空行
                self.doc.add_paragraph()
                
            i += 1
            
        # 保存文档
        self.doc.save(output_path)
        print(f'文档已保存到: {output_path}')


def main():
    """主函数"""
    # 默认值
    input_file = None
    output_file = None
    
    # 解析命令行参数
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    
    # 如果没有指定输入文件,提示用户
    if not input_file:
        print("Markdown to Word 转换工具 (增强版)")
        print("=" * 50)
        print("\n使用方法:")
        print("  python md2docx.py <输入markdown文件路径> [输出docx文件路径]")
        print("\n示例:")
        print("  python md2docx.py input.md")
        print("  python md2docx.py input.md output.docx")
        print("  python md2docx.py input.md out/my_document.docx")
        print("\n增强功能:")
        print("  ✓ 支持相对路径图片 (相对于MD文件所在目录)")
        print("  ✓ 图片自动调整大小适配页面")
        print("  ✓ 大表格自动分页并添加'续表'标题")
        print("  ✓ 一级和二级标题后自动分页")
        return
    
    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"错误: 找不到文件 '{input_file}'")
        return
    
    # 如果没有指定输出文件,使用默认路径
    if not output_file:
        # 创建out目录
        out_dir = 'out'
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)
            print(f"已创建输出目录: {out_dir}/")
        
        # 生成输出文件名
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(out_dir, f'{base_name}.docx')
    else:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"已创建输出目录: {output_dir}/")
    
    # 读取Markdown文件
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
    except Exception as e:
        print(f"错误: 无法读取文件 '{input_file}'")
        print(f"详细信息: {str(e)}")
        return
    
    # 执行转换
    try: 
       print(f"正在转换: {input_file} -> {output_file}")
       converter = MarkdownToDocxConverter(input_file_path=input_file)
       converter.convert(markdown_content, output_file)
       print("转换完成!")
       print(f"\n统计信息:")
       print(f"  - 图片数量: {converter.image_counter}")
       print(f"  - 表格数量: {converter.table_counter}")
    except Exception as e:
       print(f"转换过程中出现错误: {str(e)}")
       import traceback
       traceback.print_exc()


if __name__ == '__main__':
   main()
